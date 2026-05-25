"""
Word (.docx) report generators.

Two reports are produced:

1. main_report:    full narrative with all statistics, tables and figures.
                   Every sentence is assembled from the data — there is no
                   hardcoded biological interpretation. Optional user-
                   supplied text from data/project.json (title, subtitle,
                   sample source, extraction protocol, condition
                   descriptions, notes) is embedded if present, but the
                   report works fully without it.

2. unique_table:   focused side-by-side unique-metabolite table that
                   adapts to any number of conditions.

Both reports adapt dynamically to the conditions, groups, days,
fractions and metabolites present in the data.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ===========================================================================
# Helper utilities
# ===========================================================================
def _shade(cell, hex_color: str) -> None:
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def _add_para(doc: Document, text: str, *, size: int = 11,
              italic: bool = False, bold: bool = False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold


def _add_image(doc: Document, path: Path, width_cm: float = 16.0,
               caption: str | None = None) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)


def _add_dataframe(doc: Document, df: pd.DataFrame,
                   max_rows: int | None = None,
                   index_label: str = "") -> None:
    if df is None or df.empty:
        _add_para(doc, "(no rows to display)", italic=True, size=9)
        return
    if max_rows is not None:
        df = df.head(max_rows)
    df = df.copy()
    if df.index.name is None and index_label:
        df.index.name = index_label
    headers = [df.index.name or ""] + list(df.columns)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = str(h)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(c, "1F3A5F")
    for idx, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx)
        for j, v in enumerate(row.values, start=1):
            txt = "" if pd.isna(v) else (
                f"{v:.3f}" if isinstance(v, float) else str(v))
            cells[j].text = txt
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def _english_list(items: list[str]) -> str:
    """Format a list of strings as 'A, B, C and D'."""
    items = [str(x) for x in items if x is not None]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ", ".join(items[:-1]) + f" and {items[-1]}"


# ===========================================================================
# Data-driven narrative helpers — every sentence is built from numbers
# ===========================================================================
def _summary_stats(reid: pd.DataFrame, mat: pd.DataFrame, meta: pd.DataFrame,
                    group_order: list[str], conditions: list[str]
                    ) -> dict[str, Any]:
    n_samples = len(meta)
    n_peaks = len(reid)
    n_kept = (reid["reidentification"] == "kept_system_pick").sum() if n_peaks else 0
    n_changed = n_peaks - n_kept
    n_metab = mat.shape[0]
    n_groups = mat.shape[1]
    n_conds = len(conditions)
    n_days = len(set(meta["day"].tolist())) if not meta.empty else 0
    fractions = sorted(set(meta["fraction"].tolist())) if not meta.empty else []
    return dict(
        n_samples=n_samples, n_peaks=n_peaks, n_kept=n_kept,
        n_changed=n_changed, n_metab=n_metab, n_groups=n_groups,
        n_conds=n_conds, n_days=n_days, fractions=fractions,
        pct_changed=(100 * n_changed / max(1, n_peaks)),
    )


def _diversity_observations(div: pd.DataFrame) -> list[str]:
    """Build neutral, data-driven observations from the diversity table."""
    obs: list[str] = []
    if div.empty:
        return obs
    rich = div["richness"]
    sh = div["shannon"]
    sm = div["simpson"]
    g_max_S, g_min_S = rich.idxmax(), rich.idxmin()
    obs.append(f"Metabolite richness is highest in {g_max_S} "
               f"(S = {rich.loc[g_max_S]}) and lowest in {g_min_S} "
               f"(S = {rich.loc[g_min_S]}), a "
               f"{rich.loc[g_max_S] - rich.loc[g_min_S]}-compound spread.")
    g_max_H = sh.idxmax()
    obs.append(f"Shannon diversity (H') is highest in {g_max_H} "
               f"(H' = {sh.loc[g_max_H]:.3f}), indicating the most even "
               "distribution of metabolite abundances.")
    g_max_D = sm.idxmax()
    if g_max_D != g_max_H:
        obs.append(f"Simpson diversity (1−D) is highest in {g_max_D} "
                   f"(1−D = {sm.loc[g_max_D]:.3f}).")
    if "evenness" in div.columns:
        ev = div["evenness"]
        g_max_J = ev.idxmax()
        obs.append(f"Pielou evenness (J) is highest in {g_max_J} "
                   f"(J = {ev.loc[g_max_J]:.3f}).")
    return obs


def _per_condition_observations(buckets: dict[str, pd.DataFrame],
                                  conditions: list[str]) -> list[str]:
    obs: list[str] = []
    core_n = len(buckets.get("core_metabolites", []))
    obs.append(f"A core of {core_n} metabolites is detected in all "
               f"{len(conditions)} conditions and represents the shared "
               "metabolic backbone of the dataset.")
    counts = {c: len(buckets.get(f"unique_to_{c}", [])) for c in conditions}
    largest = max(counts, key=counts.get) if counts else None
    smallest = min(counts, key=counts.get) if counts else None
    if largest and counts[largest] > 0:
        obs.append(f"The largest condition-unique fingerprint is observed "
                   f"under {largest} ({counts[largest]} metabolites detected "
                   f"only in this condition).")
    if smallest and largest and smallest != largest:
        obs.append(f"The smallest condition-unique fingerprint is observed "
                   f"under {smallest} ({counts[smallest]} metabolites).")
    return obs


def _top_class_observations(class_rich: pd.DataFrame) -> list[str]:
    if class_rich is None or class_rich.empty:
        return []
    totals = class_rich.sum(axis=1).sort_values(ascending=False)
    top3 = totals.head(3)
    if top3.empty:
        return []
    bits = [f"{c} ({int(n)})" for c, n in top3.items()]
    return [f"The three richest compound classes across the dataset are "
            f"{_english_list(bits)} distinct metabolites respectively."]


def _top_metab_observations(mat: pd.DataFrame, top_n: int = 3) -> list[str]:
    if mat.empty:
        return []
    out = mat.copy()
    out["max"] = out.max(axis=1)
    out["argmax"] = mat.idxmax(axis=1)
    out["#groups"] = (mat > 0).sum(axis=1)
    out = out.sort_values("max", ascending=False).head(top_n)
    obs: list[str] = []
    for met, r in out.iterrows():
        if r["#groups"] >= mat.shape[1]:
            scope = f"in all {mat.shape[1]} groups"
        elif r["#groups"] == 1:
            scope = f"only in {r['argmax']}"
        else:
            scope = (f"in {int(r['#groups'])} of {mat.shape[1]} groups "
                     f"(maximum in {r['argmax']})")
        obs.append(f"{met} reaches a maximum abundance of "
                   f"{r['max']:.2f} % area, detected {scope}.")
    return obs


# ===========================================================================
# Main report
# ===========================================================================
def write_main_report(
    *,
    out_path: Path,
    figures_dir: Path,
    sample_meta: pd.DataFrame,
    parser_log: list[str],
    reid: pd.DataFrame,
    corrections: pd.DataFrame,
    mat: pd.DataFrame,
    div: pd.DataFrame,
    class_ab: pd.DataFrame,
    class_rich: pd.DataFrame,
    buckets: dict[str, pd.DataFrame],
    top40: pd.DataFrame,
    top5: pd.DataFrame,
    cond_unique: dict[str, pd.DataFrame],
    conditions: list[str],
    group_order: list[str],
    project_info: dict[str, Any] | None = None,
    quality_report: Any = None,
) -> None:
    """Generate a Word document whose narrative is fully driven by the
    data and the optional `project_info` dict."""
    info = project_info or {}
    stats = _summary_stats(reid, mat, sample_meta, group_order, conditions)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ---------- title ------------------------------------------------------
    title_text = info.get("title") or "GC-MS Metabolic Profiling Report"
    sub_text = info.get("subtitle") or (
        "Re-identification, abundance, diversity, PCA and shared/unique "
        "metabolite analysis")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(title_text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(sub_text)
    sr.italic = True
    sr.font.size = Pt(11)
    doc.add_paragraph()

    # ---------- Executive summary -----------------------------------------
    _add_heading(doc, "Executive summary", level=1)
    _add_para(doc,
        f"This report summarises GC-MS profiling of {stats['n_samples']} "
        f"sample injection(s) distributed across {stats['n_conds']} "
        f"experimental condition(s) "
        f"({_english_list(conditions) or 'unspecified'}) "
        f"and {stats['n_groups']} (condition x day) group(s). "
        f"From a total of {stats['n_peaks']} detected peak(s), "
        f"{stats['n_changed']} ({stats['pct_changed']:.1f} %) had "
        "system-assigned identifications that were either biologically "
        "implausible or library-matching artefacts and were re-identified "
        "by the candidate-walking engine using the alternative NIST hits. "
        f"After contaminant filtering and synonym normalisation, the "
        f"dataset resolves into {stats['n_metab']} distinct metabolite(s). "
        "All CSV tables and figures referenced below are available in the "
        "output directory.")

    if info.get("sample_source"):
        _add_para(doc, info["sample_source"])
    if info.get("extraction"):
        _add_para(doc, info["extraction"])

    # ---------- Sample design ---------------------------------------------
    _add_heading(doc, "1. Sample set and experimental groups", level=1)
    _add_para(doc,
        f"{stats['n_samples']} samples are organised into {stats['n_conds']} "
        f"condition(s), {stats['n_days']} distinct day value(s), and "
        f"{len(stats['fractions'])} fraction(s) "
        f"({_english_list(stats['fractions']) or 'unspecified'}). "
        "Within a (condition, day) group, all available fractions are "
        "merged by area-% summation. "
        + (info.get("fraction_note") or ""))

    # User-supplied condition descriptions (optional)
    cond_desc = info.get("conditions") or {}
    if any(cond_desc.get(c) for c in conditions):
        _add_heading(doc, "1.1 Condition descriptions", level=2)
        for c in conditions:
            txt = cond_desc.get(c)
            if txt:
                _add_para(doc, f"{c}. {txt}")

    sm = sample_meta.set_index("sample")[["condition", "day", "fraction"]]
    _add_dataframe(doc, sm, index_label="Sample")
    _add_para(doc, "Table 1 — Sample naming convention as supplied via "
                   "samples.csv.", italic=True, size=9,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------- Step 1 — parsing -----------------------------------------
    _add_heading(doc, "2. Step 1 — Parsing of Shimadzu PDF reports", level=1)
    _add_para(doc,
        "Every Shimadzu QUALITATIVE ANALYSIS REPORT contains a peak table "
        "(retention time, area, area %, height, height %, A/H ratio, "
        "system-chosen Hit-#1 compound name) followed by a library-search "
        "appendix listing up to five alternative NIST library hits per "
        "peak with similarity index (SI), molecular formula, CAS number, "
        "molecular weight and NIST retention index. The system writes only "
        "Hit #1 into the peak table regardless of how close the next hits "
        "are in similarity. The parser captures both layers so that the "
        "re-identification engine can walk the alternative hits.")
    log_rows = []
    for line in parser_log or []:
        if "peaks=" in line and "hits=" in line:
            try:
                left, peaks_part = line.split("peaks=")
                sample = left.strip()
                pn, rest = peaks_part.split("hits=")
                log_rows.append((sample, pn.strip(), rest.strip()))
            except Exception:
                continue
        elif line.startswith("FAILED"):
            log_rows.append((line.replace("FAILED", "").strip(),
                              "FAILED", ""))
    if log_rows:
        log_table = pd.DataFrame(log_rows,
                                 columns=["Sample", "Peaks", "Library hits"]
                                 ).set_index("Sample")
        _add_dataframe(doc, log_table, index_label="Sample")
        _add_para(doc, "Table 2 — Number of peaks and library hits parsed "
                       "from each PDF.", italic=True, size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------- Step 2 — re-identification --------------------------------
    _add_heading(doc, "3. Step 2 — Knowledge-based re-identification",
                 level=1)
    _add_para(doc,
        "The candidate-walking engine inspects every peak's hit list and "
        "rejects candidates that are biologically implausible for typical "
        "GC-MS biological matrices (pharmaceuticals, fungal alkaloids, "
        "scintillator dyes, halogenated synthetics, triglycerides whose "
        "retention index is incompatible with the oven program), tags "
        "lab- or instrument-derived contaminants (cyclic siloxanes, "
        "phthalates, plasticisers, perfluorinated derivatising-agent "
        "residues), and strips trimethylsilyl / TBDMS suffixes when no "
        "derivatisation step was performed. The first plausible candidate "
        "from the ordered hit list is selected and an audit trail is "
        "preserved for every peak. The full lists of substring patterns "
        "applied by the engine are configurable via gcms/config.py and "
        "can be extended for new sample matrices or analytical "
        "protocols.")
    _add_para(doc,
        f"Across all {stats['n_peaks']} peaks the system Hit-#1 pick was "
        f"retained for {stats['n_kept']} peak(s) "
        f"({100*stats['n_kept']/max(1,stats['n_peaks']):.1f} %); "
        f"{stats['n_changed']} peak(s) were re-assigned. The full audit "
        "trail (one row per peak with system_name, final_name, system_si, "
        "final_si, hit_chosen and the textual reidentification reason) is "
        "in peaks_reidentified.csv. The aggregated correction summary "
        "(distinct system_name -> final_name pairs and how often each "
        "pairing was applied) is in reidentification_corrections.csv.")

    if corrections is not None and not corrections.empty:
        _add_heading(doc, "3.1 Top 20 most-frequent corrections", level=2)
        top_corr = corrections[["system_name", "final_name", "occurrences",
                                "mean_area_pct"]].head(20).copy()
        top_corr["system_name"] = top_corr["system_name"].astype(str).str.slice(0, 50)
        top_corr["final_name"] = top_corr["final_name"].astype(str).str.slice(0, 50)
        top_corr = top_corr.rename(columns={
            "system_name": "Original system name",
            "final_name": "Re-identified as",
            "occurrences": "n",
            "mean_area_pct": "mean area %",
        }).set_index("Original system name")
        _add_dataframe(doc, top_corr, index_label="Original system name")
        _add_para(doc, "Table 3 — Top 20 most-frequent system → final "
                       "corrections.", italic=True, size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------- Step 2b — Data quality cleanup ---------------------------
    if quality_report is not None:
        _add_heading(doc, "4. Step 2b — PubChem enrichment, data quality "
                            "cleanup and deduplication", level=1)
        _add_para(doc,
            "Before any abundance / diversity statistics are computed, "
            "the re-identified peak table is passed through a quality-"
            "control pipeline that performs three things in sequence: "
            "(a) PubChem-API enrichment to obtain canonical IUPAC names "
            "and InChIKeys for every compound, (b) deduplication using "
            "the InChIKey-14 connectivity hash so that stereoisomers, "
            "salt forms and tautomers of the same compound are correctly "
            "merged, and (c) a battery of error filters that remove "
            "missing / non-positive values, solvent-front peaks, low-"
            "similarity library matches and (optionally) trace peaks. "
            "Thresholds are configurable via data/quality.json (a "
            "template called quality_template.json is auto-generated on "
            "first run).")
        _add_para(doc,
            "PubChem (https://pubchem.ncbi.nlm.nih.gov) is a free public "
            "compound database maintained by the U.S. National Library "
            "of Medicine. Every distinct compound name from the re-"
            "identified peak table is queried via the PUG REST API to "
            "retrieve the canonical IUPAC name, the PubChem CID, the "
            "InChIKey (a 27-character compound hash where the first 14 "
            "characters identify the molecular skeleton independent of "
            "stereochemistry, salts and tautomers), the molecular "
            "formula and the exact monoisotopic mass. Lookups are "
            "cached in data/pubchem_cache.json so that re-runs reuse "
            "the cached results instantly, and the tool degrades "
            "gracefully on network failure: any compound that cannot be "
            "matched simply keeps its original NIST-derived name.")
        _add_para(doc,
            "Cleaning steps applied (in order):")
        steps = [
            "PubChem name → CID → IUPACName / InChIKey / formula / "
            "exact-mass lookup; replace the NIST name with the PubChem "
            "canonical IUPAC name when one is available.",
            "Compute compound_key = first 14 characters of the InChIKey "
            "(falls back to the lower-cased final name when the "
            "InChIKey is missing).",
            "Drop rows with missing or non-positive area_pct, retention "
            "time, height or area.",
            "Drop solvent-front peaks below the configured minimum "
            "retention time (default 2.0 min — the Shimadzu method's "
            "solvent cut was at 1.5 min in the supplied data).",
            "Drop low-similarity peaks (final_si below the configured "
            "threshold, default 50). Below this cut-off the library "
            "match is too weak to trust.",
            "Drop trace peaks (area_pct below the configured threshold; "
            "default 0 = off, but can be set to e.g. 0.05 to remove "
            "integration-noise peaks).",
            "Normalise compound-name spelling and stereochemistry "
            "variants on the rare compounds that PubChem could not "
            "match (`Oleic Acid, (Z)-` → `Oleic acid`, `n-Hexadecanoic "
            "acid` → `Hexadecanoic acid`, etc.).",
            "Merge within-sample co-elutions on compound_key: when the "
            "same compound appears in multiple peaks of the same sample "
            "(broadened peaks, shoulders, integration splits, "
            "stereoisomers detected as separate peaks) the rows are "
            "summed into one (sum of area / area_pct, max similarity, "
            "median retention time).",
            "Optionally drop dataset-wide singletons (compounds detected "
            "in only one sample at very low abundance). Off by default.",
        ]
        for s in steps:
            _add_para(doc, "• " + s)

        _add_dataframe(doc, quality_report.as_dataframe(),
                       index_label="Quality metric")
        _add_para(doc,
            "Table 4 — Number of rows removed at each quality-control "
            "step. The full audit trails are in pubchem_enrichment.csv "
            "(name → CID → InChIKey → formula → exact mass for every "
            "compound), pubchem_name_remap.csv (NIST → PubChem name "
            "renames), synonym_collapses.csv and quality_report.csv.",
            italic=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        coverage = (1 - quality_report.n_out / max(1, quality_report.n_in)) * 100
        _add_para(doc,
            f"In total, {quality_report.n_in - quality_report.n_out} of "
            f"{quality_report.n_in} re-identified rows ({coverage:.1f} %) "
            "were either removed as errors / noise or merged as "
            "co-elutions. The cleaned table contains "
            f"{quality_report.n_out} rows and is the input to every "
            "down-stream statistic in this report. The cleaned table is "
            "saved as peaks_cleaned.csv.")

    # ---------- Step 3 — abundance matrix --------------------------------
    _add_heading(doc, "5. Step 3 — Metabolite x group abundance matrix",
                 level=1)
    _add_para(doc,
        "Within each (condition, day) group the area-% values from all "
        "available fractions are summed. Compounds tagged as Unidentified "
        "or Contaminant are excluded so the down-stream diversity "
        "statistics reflect only genuine metabolites. A light synonym "
        "normalisation step then collapses obvious aliases.")
    _add_para(doc,
        f"The resulting matrix has {mat.shape[0]} metabolite(s) and "
        f"{mat.shape[1]} group(s). It contains "
        f"{int((mat > 0).sum().sum())} non-zero cells "
        f"({100*(mat > 0).sum().sum()/max(1, mat.shape[0]*mat.shape[1]):.1f} %"
        " density). The full matrix is saved as "
        "metabolite_abundance_matrix.csv; the presence / absence view "
        "with the #groups column is saved as metabolite_presence.csv.")

    # ---------- Step 4 — alpha-diversity ---------------------------------
    _add_heading(doc, "6. Step 4 — Alpha-diversity per group", level=1)
    _add_para(doc,
        "Four indices are computed per group from the metabolite × group "
        "matrix. Richness (S) is the count of distinct metabolites with "
        "non-zero signal. Shannon entropy H' = − Σ p_i × ln(p_i), where "
        "p_i is the relative abundance of metabolite i within the group. "
        "Gini−Simpson 1−D = 1 − Σ p_i² reflects the probability that two "
        "randomly drawn signal units belong to different metabolites. "
        "Pielou evenness J = H' / ln(S) decouples evenness from richness. "
        "Total signal is the sum of area-% across all metabolites in the "
        "group.")
    _add_dataframe(doc, div, index_label="Group")
    _add_para(doc, "Table 5 — Richness, Shannon, Simpson and evenness per "
                   "group.", italic=True, size=9,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_image(doc, figures_dir / "03_alpha_diversity.png", width_cm=16.0,
               caption="Figure 1 — Alpha-diversity indices per group.")

    obs = _diversity_observations(div)
    if obs:
        _add_heading(doc, "6.1 Observations from the diversity statistics",
                     level=2)
        for o in obs:
            _add_para(doc, "• " + o)

    # ---------- Step 5 — PCA ---------------------------------------------
    _add_heading(doc, "7. Step 5 — Principal-component ordination", level=1)
    _add_para(doc,
        "Abundance values are log10(1 + x) transformed to dampen the "
        "influence of the most abundant peaks, transposed so rows are "
        "groups and columns metabolites, standardised column-wise (zero "
        "mean, unit variance) and projected onto the first two principal "
        "components. The proportion of total variance captured by PC1 and "
        "PC2 is shown on the figure axes. Each group is plotted as a "
        "single point coloured by group identity.")
    _add_image(doc, figures_dir / "01_PCA.png", width_cm=15.0,
               caption="Figure 2 — PCA on log-transformed metabolite "
                       "abundances.")

    # ---------- Step 6 — class diversity ---------------------------------
    _add_heading(doc, "8. Step 6 — Compound-class diversity", level=1)
    classes_present = (sorted(class_rich.index.tolist())
                        if class_rich is not None and not class_rich.empty
                        else [])
    if classes_present:
        _add_para(doc,
            "Each re-identified compound is assigned to a biochemical "
            "class by the regular-expression classifier in gcms/config.py "
            "(rule-ordered: more specific patterns first). In this "
            f"dataset {len(classes_present)} class(es) are populated: "
            f"{_english_list(classes_present)}. The class table is "
            "summarised in two ways: by abundance (sum area-% per class "
            "per group) and by richness (distinct compounds per class "
            "per group).")
    else:
        _add_para(doc,
            "Each re-identified compound is assigned to a biochemical "
            "class by the regular-expression classifier in gcms/config.py. "
            "No populated classes were produced from this dataset.")
    _add_image(doc, figures_dir / "02_class_diversity_bar.png", width_cm=16.0,
               caption="Figure 3 — Stacked bar of compound-class richness.")
    _add_image(doc, figures_dir / "05_heatmap_classes.png", width_cm=15.5,
               caption="Figure 4 — Heat-map of compound-class abundance.")
    _add_dataframe(doc, class_rich.fillna(0).astype(int)
                   if class_rich is not None else pd.DataFrame(),
                   index_label="Compound class")
    _add_para(doc, "Table 6 — Number of distinct metabolites per class per "
                   "group.", italic=True, size=9,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    cls_obs = _top_class_observations(class_rich)
    for o in cls_obs:
        _add_para(doc, "• " + o)

    # ---------- Step 7 — top metabolites heatmap -------------------------
    _add_heading(doc, "9. Step 7 — Heat-map of top metabolites", level=1)
    _add_para(doc,
        "The top 40 metabolites by maximum across-group abundance are "
        "plotted as a log10(area-% + 0.01) heat-map. Choosing the top by "
        "maximum (rather than mean) preserves rows that are highly "
        "abundant in a single condition but absent in others — exactly "
        "the kind of stress / condition-specific marker that the "
        "analysis is designed to highlight.")
    _add_image(doc, figures_dir / "04_heatmap_top40_metabolites.png",
               width_cm=15.0,
               caption="Figure 5 — Top-40 metabolites by max across-group "
                       "abundance (log10 area-%).")
    for o in _top_metab_observations(mat):
        _add_para(doc, "• " + o)

    # ---------- Step 8 — shared / unique ---------------------------------
    _add_heading(doc, "10. Step 8 — Shared vs unique metabolites", level=1)
    _add_image(doc, figures_dir / "06_unique_vs_core.png", width_cm=14.0,
               caption="Figure 6 — Number of metabolites unique to each "
                       "condition versus the shared core.")
    for o in _per_condition_observations(buckets, conditions):
        _add_para(doc, "• " + o)

    # ---------- Annexe A — unique to each condition ----------------------
    doc.add_page_break()
    _add_heading(doc, "Annexe A. Metabolites unique to each condition",
                 level=1)
    _add_para(doc,
        "Each list contains metabolites detected in only one of the "
        f"{len(conditions)} condition(s) (all fractions and all days "
        "merged), and absent in every other condition.")
    for c in conditions:
        df = cond_unique.get(c, pd.DataFrame())
        _add_heading(doc, f"A.{c} — Unique to {c} ({len(df)} compounds)",
                     level=2)
        if df is None or df.empty:
            _add_para(doc, "(none)")
        else:
            _add_dataframe(doc, df, index_label="Metabolite")

    # ---------- Annexe B — most abundant overall -------------------------
    doc.add_page_break()
    _add_heading(doc, "Annexe B. Most abundant metabolites overall", level=1)
    _add_para(doc,
        "Top-40 metabolites ranked by the maximum across-group abundance. "
        "max_area_pct = highest sum-of-area-% in any single group; "
        "mean_area_pct = average across all groups; "
        "argmax_group = group where the maximum was observed. A "
        "metabolite that reaches a high max but appears in only a few "
        "groups is a strong condition marker; a metabolite with a high "
        "mean and #groups_present equal to the total number of groups "
        "is part of the constitutive core.")
    _add_dataframe(doc, top40, index_label="Metabolite")

    # ---------- Annexe C — top-5 per group ------------------------------
    doc.add_page_break()
    _add_heading(doc, "Annexe C. Top-5 most abundant metabolites in each "
                       "group", level=1)
    if top5 is not None and not top5.empty:
        for g in group_order:
            sub = top5[top5["group"] == g].drop(columns=["group"]).set_index("rank")
            if sub.empty:
                continue
            _add_heading(doc, f"C.{g}", level=2)
            _add_dataframe(doc, sub, index_label="Rank")

    # ---------- Annexe D — files produced --------------------------------
    doc.add_page_break()
    _add_heading(doc, "Annexe D. Files produced", level=1)
    _add_para(doc,
        "Every artefact referenced in this report is available under "
        "the output/ directory of the project.")
    files_table = pd.DataFrame([
        ("peaks_raw.csv",                       "Raw parsed peak table per sample"),
        ("hits_raw.csv",                        "Raw NIST library hits per peak"),
        ("peaks_reidentified.csv",              "Re-identified peaks (audit trail)"),
        ("reidentification_corrections.csv",    "Distinct system → final corrections"),
        ("pubchem_enrichment.csv",              "NIST name → PubChem CID/IUPAC/InChIKey/formula/mass"),
        ("pubchem_name_remap.csv",              "NIST → PubChem canonical name remaps"),
        ("data/pubchem_cache.json",             "PubChem lookup cache (re-run instantly)"),
        ("synonym_collapses.csv",               "NIST → normalised name collapses"),
        ("quality_report.csv",                  "Per-step row counts of the QC pipeline"),
        ("peaks_cleaned.csv",                   "Cleaned peak table after QC"),
        ("metabolite_abundance_matrix.csv",     "Metabolite × group abundance"),
        ("metabolite_presence.csv",             "Presence / absence + #groups"),
        ("diversity_indices.csv",               "Richness, Shannon, Simpson, evenness"),
        ("class_abundance.csv",                 "Compound class × group abundance"),
        ("class_richness.csv",                  "Compound class × group richness"),
        ("set_core_metabolites.csv",            "Core (in all conditions)"),
        ("set_unique_to_<cond>.csv",            "Unique to each condition"),
        ("lists/top40_most_abundant.csv",       "Top 40 by max across-group abundance"),
        ("lists/top5_per_group.csv",            "Top 5 per group"),
        ("lists/unique_<cond>.csv",             "Per-condition unique tables"),
        ("figures/01_PCA.png",                  "PCA"),
        ("figures/02_class_diversity_bar.png",  "Class diversity bar"),
        ("figures/03_alpha_diversity.png",      "Alpha-diversity"),
        ("figures/04_heatmap_top40_metabolites.png", "Top-40 metabolite heatmap"),
        ("figures/05_heatmap_classes.png",      "Class abundance heatmap"),
        ("figures/06_unique_vs_core.png",       "Unique vs core counts"),
    ], columns=["File", "Description"]).set_index("File")
    _add_dataframe(doc, files_table, index_label="File")

    if info.get("notes"):
        doc.add_page_break()
        _add_heading(doc, "Notes", level=1)
        _add_para(doc, info["notes"])

    doc.save(str(out_path))


# ===========================================================================
# Standalone unique-metabolites table
# ===========================================================================
def write_unique_table(
    *,
    out_path: Path,
    cond_unique: dict[str, pd.DataFrame],
    conditions: list[str],
    project_info: dict[str, Any] | None = None,
) -> None:
    info = project_info or {}
    palette_hex = ["2CA02C", "D62728", "9467BD", "1F77B4",
                   "FF7F0E", "17BECF", "BCBD22", "E377C2"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(info.get("unique_table_title")
                      or "Metabolites unique to each condition")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(" · ".join(conditions)
                     + " (all fractions and all days merged)")
    sr.italic = True
    sr.font.size = Pt(11)
    doc.add_paragraph()

    n_rows = max((len(cond_unique.get(c, pd.DataFrame()))
                  for c in conditions), default=0)
    if n_rows == 0:
        _add_para(doc, "No metabolites were uniquely detected in any "
                       "single condition.")
        doc.save(str(out_path))
        return

    n_cols = 2 * len(conditions)
    table = doc.add_table(rows=2, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr1 = table.rows[0].cells
    for i, c in enumerate(conditions):
        a, b = 2 * i, 2 * i + 1
        hdr1[a].merge(hdr1[b])
        n = len(cond_unique.get(c, pd.DataFrame()))
        hdr1[a].text = f"Unique to {c}  (n = {n})"
        for p in hdr1[a].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr1[a], palette_hex[i % len(palette_hex)])

    hdr2 = table.rows[1].cells
    for i in range(len(conditions)):
        hdr2[2 * i].text = "Metabolite"
        hdr2[2 * i + 1].text = "Area %"
        for j in (2 * i, 2 * i + 1):
            for p in hdr2[j].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(hdr2[j], "1F3A5F")

    for i in range(n_rows):
        row = table.add_row().cells
        for k, c in enumerate(conditions):
            df = cond_unique.get(c, pd.DataFrame()).reset_index()
            if i < len(df):
                row[2 * k].text = str(df.iloc[i]["metabolite"])
                row[2 * k + 1].text = f"{df.iloc[i]['abundance_area_pct']:.3f}"
            for j in (2 * k, 2 * k + 1):
                for p in row[j].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Side-by-side comparison of metabolites detected "
                     "exclusively under one condition.")
    cr.italic = True
    cr.font.size = Pt(9)

    # Per-condition individual tables
    for i, c in enumerate(conditions):
        df = cond_unique.get(c, pd.DataFrame())
        if df.empty:
            continue
        doc.add_page_break()
        _add_heading(doc, f"Metabolites unique to {c}  (n = {len(df)})",
                     level=1)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        h = t.rows[0].cells
        h[0].text = "Metabolite"
        h[1].text = "Abundance (sum area %)"
        for cell in h:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(cell, palette_hex[i % len(palette_hex)])
        for met, row in df.iterrows():
            cells = t.add_row().cells
            cells[0].text = str(met)
            cells[1].text = f"{row['abundance_area_pct']:.3f}"
            for cell in cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

    doc.save(str(out_path))
